import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, timedelta
from database import DatabaseManager
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import os


class ProfitReportTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.db = DatabaseManager()
        self.setup_ui()
        self.filtered_items = []
        self.update_report()

    def setup_ui(self):
        # Main container with modern styling
        main_container = ttk.Frame(self)
        main_container.pack(fill="both", expand=True, padx=20, pady=20)

        # Header section
        header_frame = ttk.Frame(main_container)
        header_frame.pack(fill="x", pady=(0, 20))

        ttk.Label(header_frame, text="📊 Profit Analytics",
                  font=("Segoe UI", 24, "bold")).pack(side="left")

        export_btn = ttk.Button(header_frame, text="📄 Export Report",
                                command=self.export_to_word, style="Accent.TButton")
        export_btn.pack(side="right")

        # Controls section
        controls_frame = ttk.LabelFrame(main_container, text="Report Configuration", padding=15)
        controls_frame.pack(fill="x", pady=(0, 20))

        # First row of controls
        row1 = ttk.Frame(controls_frame)
        row1.pack(fill="x", pady=(0, 10))

        ttk.Label(row1, text="Period:", font=("Segoe UI", 10)).pack(side="left")
        self.report_type = tk.StringVar(value="Day")
        period_combo = ttk.Combobox(row1, textvariable=self.report_type,
                                    values=["Day", "Week", "Month", "Year"],
                                    state="readonly", width=12)
        period_combo.pack(side="left", padx=(10, 30))
        period_combo.bind("<<ComboboxSelected>>", self.update_report)

        ttk.Label(row1, text="Date:", font=("Segoe UI", 10)).pack(side="left")
        self.date_entry = ttk.Entry(row1, width=15, font=("Segoe UI", 10))
        self.date_entry.insert(0, datetime.now().strftime('%Y-%m-%d'))
        self.date_entry.pack(side="left", padx=(10, 0))
        self.date_entry.bind("<KeyRelease>", self.update_report)

        # Metrics display
        self.metrics_frame = ttk.LabelFrame(main_container, text="Key Metrics", padding=20)
        self.metrics_frame.pack(fill="both", expand=True)

    def update_report(self, event=None):
        # Clear existing metrics
        for widget in self.metrics_frame.winfo_children():
            widget.destroy()

        try:
            selected_date = datetime.strptime(self.date_entry.get(), '%Y-%m-%d')
        except ValueError:
            ttk.Label(self.metrics_frame, text="⚠️ Invalid date format. Use YYYY-MM-DD",
                      foreground="red").pack()
            return

        # Calculate date range
        report_type = self.report_type.get()
        if report_type == "Day":
            start_date = selected_date
        elif report_type == "Week":
            start_date = selected_date - timedelta(days=selected_date.weekday())
        elif report_type == "Month":
            start_date = selected_date.replace(day=1)
        else:  # Year
            start_date = selected_date.replace(month=1, day=1)

        # Process data
        self.filtered_items = []
        metrics = {"cost": 0, "revenue": 0, "profit": 0, "count": 0}

        for name, source, sold, date_str in self.db.fetch_items():
            try:
                if sold > 0 and datetime.strptime(date_str, "%Y-%m-%d") >= start_date:
                    self.filtered_items.append((name, source, sold, date_str))
                    metrics["cost"] += source
                    metrics["revenue"] += sold
                    metrics["profit"] += (sold - source)
                    metrics["count"] += 1
            except:
                continue

        # Create metrics grid
        grid_frame = ttk.Frame(self.metrics_frame)
        grid_frame.pack(fill="both", expand=True)

        metric_cards = [
            ("📦 Items Sold", f"{metrics['count']}", "items"),
            ("💰 Revenue", f"₱{metrics['revenue']:,.2f}", "total sales"),
            ("💸 Cost", f"₱{metrics['cost']:,.2f}", "total expenses"),
            ("📈 Profit", f"₱{metrics['profit']:,.2f}", "net gain"),
        ]

        for i, (title, value, subtitle) in enumerate(metric_cards):
            card = ttk.LabelFrame(grid_frame, text=title, padding=15)
            card.grid(row=i // 2, column=i % 2, padx=10, pady=10, sticky="nsew")

            ttk.Label(card, text=value, font=("Segoe UI", 16, "bold")).pack()
            ttk.Label(card, text=subtitle, font=("Segoe UI", 9),
                      foreground="gray").pack()

        # Configure grid weights
        grid_frame.columnconfigure(0, weight=1)
        grid_frame.columnconfigure(1, weight=1)
        grid_frame.rowconfigure(0, weight=1)
        grid_frame.rowconfigure(1, weight=1)

        # Add profit margin
        if metrics["revenue"] > 0:
            margin = (metrics["profit"] / metrics["revenue"]) * 100
            margin_label = ttk.Label(self.metrics_frame,
                                     text=f"📊 Profit Margin: {margin:.1f}%",
                                     font=("Segoe UI", 12, "bold"))
            margin_label.pack(pady=(20, 0))

    def create_profit_chart(self):
        """Create a profit breakdown chart"""
        if not self.filtered_items:
            return None

        # Aggregate data by date
        daily_profits = {}
        for name, source, sold, date_str in self.filtered_items:
            profit = sold - source
            if date_str in daily_profits:
                daily_profits[date_str] += profit
            else:
                daily_profits[date_str] = profit

        # Create chart
        plt.figure(figsize=(10, 6))
        dates = list(daily_profits.keys())
        profits = list(daily_profits.values())

        plt.bar(dates, profits, color='#2E86AB', alpha=0.8)
        plt.title('Daily Profit Analysis', fontsize=16, fontweight='bold')
        plt.xlabel('Date', fontsize=12)
        plt.ylabel('Profit (₱)', fontsize=12)
        plt.xticks(rotation=45)
        plt.grid(axis='y', alpha=0.3)
        plt.tight_layout()

        # Save to bytes
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()

        return img_buffer

    def export_to_word(self):
        if not self.filtered_items:
            messagebox.showwarning("No Data", "No data available for export.")
            return

        doc = Document()

        # Header
        header = doc.add_heading('PROFIT ANALYSIS REPORT', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f"Period: {self.report_type.get()} | Generated: {datetime.now().strftime('%B %d, %Y')}")

        # Calculate totals
        totals = {
            'cost': sum(source for _, source, _, _ in self.filtered_items),
            'revenue': sum(sold for _, _, sold, _ in self.filtered_items),
        }
        totals['profit'] = totals['revenue'] - totals['cost']
        totals['margin'] = (totals['profit'] / totals['revenue'] * 100) if totals['revenue'] > 0 else 0

        # Summary section
        summary = doc.add_paragraph()
        summary.add_run("EXECUTIVE SUMMARY\n").bold = True
        summary.add_run(f"Items Sold: {len(self.filtered_items)} | ")
        summary.add_run(f"Revenue: ₱{totals['revenue']:,.2f} | ")
        summary.add_run(f"Profit: ₱{totals['profit']:,.2f} | ")
        summary.add_run(f"Margin: {totals['margin']:.1f}%")

        # Add chart
        chart_buffer = self.create_profit_chart()
        if chart_buffer:
            doc.add_paragraph("\nPROFIT TREND ANALYSIS").bold = True
            chart_para = doc.add_paragraph()
            run = chart_para.runs[0] if chart_para.runs else chart_para.add_run()
            run.add_picture(chart_buffer, width=Inches(6))

        # Transaction table
        doc.add_paragraph("\nDETAILED TRANSACTIONS").bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        headers = ['Item', 'Cost', 'Revenue', 'Profit', 'Date']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for name, source, sold, date in self.filtered_items:
            row = table.add_row()
            cells = row.cells
            cells[0].text = name
            cells[1].text = f"₱{source:,.2f}"
            cells[2].text = f"₱{sold:,.2f}"
            cells[3].text = f"₱{sold - source:,.2f}"
            cells[4].text = date

        # Save
        filename = f"Profit_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        messagebox.showinfo("Success", f"Report exported: {filename}")